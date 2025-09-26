import streamlit as st
import tempfile
import os
import sys
from typing import List, Optional
from dataclasses import dataclass
from datetime import datetime
import base64

# Document processing
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document

# Simple LLM integration (choose one)
# Option 1: OpenAI
# from openai import OpenAI
# Option 2: Groq
from langchain_groq import ChatGroq
# Option 3: Ollama (local)
# from langchain_ollama import OllamaLLM

# Load API key
try:
    api_key = st.secrets.get("GROQ_API_KEY") or os.getenv("GROQ_API_KEY")
except:
    api_key = os.getenv("GROQ_API_KEY")

if not api_key:
    st.error("⚠️ No GROQ_API_KEY found. Please set it in Streamlit Secrets or .env")
    st.stop()

# Configure page
st.set_page_config(
    page_title="Quick Query",
    page_icon="📚",
    layout="centered",
    initial_sidebar_state="expanded"
)

@dataclass
class ChatMessage:
    """Simple chat message class"""
    role: str  # 'user' or 'assistant'
    content: str
    timestamp: datetime
    sources: Optional[List[str]] = None

class SimpleDocumentQA:
    """Simplified document Q&A system"""
    
    def __init__(self):
        self.vectorstore = None
        self.documents_metadata = []
        
        # Text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        
        # Embeddings
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        # LLM
        try:
            self.llm = ChatGroq(
                groq_api_key=api_key,
                model="llama-3.1-8b-instant",
                temperature=0.7,
                max_tokens=1500
            )
        except Exception as e:
            st.error(f"Failed to initialize LLM: {e}")
            self.llm = None
    
    def extract_text_from_file(self, uploaded_file) -> tuple[str, str]:
        """Extract text from uploaded file"""
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=uploaded_file.name) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        filename = uploaded_file.name
        file_ext = filename.lower().split('.')[-1]
        
        try:
            if file_ext == 'pdf':
                text = self._extract_pdf_text(tmp_path)
            elif file_ext == 'docx':
                text = self._extract_docx_text(tmp_path)
            elif file_ext == 'txt':
                text = self._extract_txt_text(tmp_path)
            else:
                st.error(f"Unsupported file type: {file_ext}")
                text = ""
            
            return filename, text
        finally:
            os.unlink(tmp_path)
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            st.error(f"Error processing PDF: {str(e)}")
            return ""
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error processing DOCX: {str(e)}")
            return ""
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            st.error(f"Error processing TXT: {str(e)}")
            return ""
    
    def load_documents(self, uploaded_files) -> bool:
        """Process and store documents"""
        if not uploaded_files:
            return False
        
        documents = []
        progress_bar = st.progress(0)
        
        for i, uploaded_file in enumerate(uploaded_files):
            filename, text = self.extract_text_from_file(uploaded_file)
            
            if text.strip():
                doc = Document(
                    page_content=text,
                    metadata={"source": filename, "upload_time": datetime.now().isoformat()}
                )
                documents.append(doc)
            
            progress_bar.progress((i + 1) / len(uploaded_files))
        
        if documents:
            return self._create_vectorstore(documents)
        return False
    
    def _create_vectorstore(self, documents: List[Document]) -> bool:
        """Create FAISS vectorstore from documents"""
        try:
            # Split documents
            chunks = self.text_splitter.split_documents(documents)
            
            if not chunks:
                st.error("No text chunks created")
                return False
            
            # Create or update vectorstore
            if self.vectorstore is None:
                self.vectorstore = FAISS.from_documents(chunks, self.embeddings)
                self.documents_metadata = [doc.metadata for doc in chunks]
            else:
                new_vectorstore = FAISS.from_documents(chunks, self.embeddings)
                self.vectorstore.merge_from(new_vectorstore)
                self.documents_metadata.extend([doc.metadata for doc in chunks])
            
            # Store in session state
            st.session_state['vectorstore'] = self.vectorstore
            st.session_state['docs_metadata'] = self.documents_metadata
            
            return True
        except Exception as e:
            st.error(f"Error creating vectorstore: {str(e)}")
            return False
    
    def search_documents(self, query: str, k: int = 5) -> List[Document]:
        """Search for relevant documents"""
        if self.vectorstore:
            try:
                return self.vectorstore.similarity_search(query, k=k)
            except Exception as e:
                st.error(f"Search error: {e}")
                return []
        return []
    
    def generate_answer(self, query: str) -> tuple[str, List[str]]:
        """Generate answer using retrieved documents"""
        if not self.llm:
            return "LLM not available", []
        
        # Search for relevant documents
        relevant_docs = self.search_documents(query, k=3)
        
        if not relevant_docs:
            return "I couldn't find relevant information in your documents to answer this question.", []
        
        # Prepare context
        context = "\n\n".join([doc.page_content for doc in relevant_docs])
        sources = [doc.metadata.get('source', 'Unknown') for doc in relevant_docs]
        
        # Create prompt
        prompt = f"""Based on the following context from uploaded documents, answer the user's question. 
If the information is not in the context, say so clearly.

Context:
{context}

Question: {query}

Answer:"""
        
        try:
            # Generate response
            response = self.llm.invoke(prompt)
            return response.content, list(set(sources))  # Remove duplicate sources
        except Exception as e:
            return f"Error generating response: {str(e)}", []
    
    def clear_documents(self) -> bool:
        """Clear all documents"""
        try:
            self.vectorstore = None
            self.documents_metadata = []
            
            if 'vectorstore' in st.session_state:
                del st.session_state['vectorstore']
            if 'docs_metadata' in st.session_state:
                del st.session_state['docs_metadata']
            
            return True
        except Exception as e:
            st.error(f"Error clearing documents: {e}")
            return False
    
    def get_document_count(self) -> int:
        """Get number of document chunks"""
        return len(self.documents_metadata)

def load_image_as_base64(image_path: str) -> str:
    """Load image and convert to base64"""
    try:
        if os.path.exists(image_path):
            with open(image_path, "rb") as img_file:
                return base64.b64encode(img_file.read()).decode()
    except Exception:
        pass
    return ""

def render_sidebar():
    """Render sidebar with controls"""
    with st.sidebar:
        # Logo
        logo_path = "images/Quickquery.png"
        logo_b64 = load_image_as_base64(logo_path)
        
        if logo_b64:
            st.markdown(f"""
                <div style="text-align: center; margin-bottom: 20px;">
                    <img src="data:image/png;base64,{logo_b64}" 
                        style="width: 140px; height: auto; margin-bottom: 10px;" />
                    <p style="margin: 0; font-style: italic; color: #666; font-size: 14px;">
                        Simple Document Q&A
                    </p>
                </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("# 📚 Quick Query")
            st.markdown("*Simple Document Q&A*")
        
        st.markdown("---")
        
        # Status
        status = "✅ Connected" if api_key else "❌ Not Found"
        st.markdown(f"**API Status:** {status}")
        
        st.markdown("---")
        
        # Document management
        st.markdown("### 📁 Document Management")
        
        doc_count = st.session_state.get('doc_count', 0)
        if doc_count > 0:
            st.success(f"✅ {doc_count} document chunks loaded")
            
            if st.button("🗑️ Clear All Documents"):
                if st.session_state.qa_system.clear_documents():
                    st.session_state['documents_loaded'] = False
                    st.session_state['doc_count'] = 0
                    st.success("Documents cleared!")
                    st.rerun()
        
        # File uploader
        uploaded_files = st.file_uploader(
            "Upload Documents",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            help="Upload PDF, DOCX, or TXT files"
        )
        
        if uploaded_files:
            if st.button("📤 Process Documents"):
                with st.spinner("Processing documents..."):
                    success = st.session_state.qa_system.load_documents(uploaded_files)
                    if success:
                        st.session_state['documents_loaded'] = True
                        st.session_state['doc_count'] = st.session_state.qa_system.get_document_count()
                        st.success("Documents processed successfully!")
                        st.rerun()
                    else:
                        st.error("Failed to process documents")
        
        return uploaded_files

def render_chat_message(message: ChatMessage, show_sources: bool = True):
    """Render a chat message"""
    if message.role == "user":
        st.markdown(f"""
        <div style="background-color: #e3f2fd; padding: 15px; border-radius: 10px; margin: 10px 0; text-align: right;">
            <strong>You:</strong><br>
            {message.content}
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div style="background-color: #f5f5f5; padding: 15px; border-radius: 10px; margin: 10px 0; border-left: 4px solid #2196f3;">
            <strong>Quick Query:</strong><br>
            {message.content}
        </div>
        """, unsafe_allow_html=True)
        
        if show_sources and message.sources:
            st.markdown("**📚 Sources:**")
            for source in message.sources:
                st.markdown(f"• {source}")

def main():
    """Main application"""
    
    # Initialize session state
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'qa_system' not in st.session_state:
        st.session_state.qa_system = SimpleDocumentQA()
    if 'documents_loaded' not in st.session_state:
        st.session_state.documents_loaded = False
    if 'doc_count' not in st.session_state:
        st.session_state.doc_count = 0
    
    # Load existing vectorstore if available
    if 'vectorstore' in st.session_state:
        st.session_state.qa_system.vectorstore = st.session_state['vectorstore']
        st.session_state.qa_system.documents_metadata = st.session_state.get('docs_metadata', [])
        st.session_state['documents_loaded'] = True
        st.session_state['doc_count'] = len(st.session_state.qa_system.documents_metadata)
    
    # CSS
    st.markdown("""
    <style>
    .main-header {
        text-align: left;
        padding: 20px;
        margin-bottom: 30px;
        background: white;
        color: #333;
        border-radius: 10px;
        border: 1px solid #e0e0e0;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Render sidebar
    uploaded_files = render_sidebar()
    
    # Main header
    header_logo_path = "images/Quickquery.png"
    header_logo_b64 = load_image_as_base64(header_logo_path)
    
    if header_logo_b64:
        st.markdown(f'''
        <div class="main-header">
            <div style="display: flex; align-items: flex-end;">
                <img src="data:image/png;base64,{header_logo_b64}" 
                    style="width: 240px; height: auto; margin-right: 20px;" />
                <h1 style="margin: 0; font-size: 28px; font-weight: bold;">
                    Find Answers Inside Your Documents
                </h1>
            </div>
        </div>
        ''', unsafe_allow_html=True)
    else:
        st.markdown('''
        <div class="main-header">
            <h1>📚 Find Answers Inside Your Documents</h1>
        </div>
        ''', unsafe_allow_html=True)
    
    # Dynamic topics section based on uploaded documents
    if st.session_state.documents_loaded and st.session_state.doc_count > 0:
        st.markdown("---")
        st.markdown("### 📚 Available Documents")
        
        # Get unique document sources from metadata
        if hasattr(st.session_state.qa_system, 'documents_metadata') and st.session_state.qa_system.documents_metadata:
            unique_sources = list(set([
                doc_meta.get('source', 'Unknown Document') 
                for doc_meta in st.session_state.qa_system.documents_metadata
            ]))
            
            # Display documents as a nice list
            for source in sorted(unique_sources):
                # Clean up the filename for display
                display_name = source.replace('.pdf', '').replace('.docx', '').replace('.txt', '')
                display_name = display_name.replace('_', ' ').replace('-', ' ').title()
                st.markdown(f"• **{display_name}** ({source})")
        
        st.markdown("*Ask questions about any of these documents!*")
    else:
        st.markdown("---")
        st.markdown("### 📚 Upload Documents to Get Started")
        st.markdown("Once you upload documents, they will appear here and you can ask questions about them.")
    
    # Chat interface
    for message in st.session_state.chat_history:
        render_chat_message(message, show_sources=True)
    
    # Input area
    if st.session_state.documents_loaded:
        user_input = st.chat_input("Ask a question about your documents...")
        
        if user_input:
            # Add user message
            user_message = ChatMessage(
                role="user",
                content=user_input,
                timestamp=datetime.now()
            )
            st.session_state.chat_history.append(user_message)
            
            # Generate answer
            with st.spinner("Analyzing your question..."):
                answer, sources = st.session_state.qa_system.generate_answer(user_input)
            
            # Add assistant message
            assistant_message = ChatMessage(
                role="assistant",
                content=answer,
                timestamp=datetime.now(),
                sources=sources
            )
            st.session_state.chat_history.append(assistant_message)
            
            st.rerun()
    else:
        st.info("👆 Please upload documents using the sidebar to start asking questions!")
    
    # Clear chat button
    if st.session_state.chat_history:
        if st.button("🗑️ Clear Chat History"):
            st.session_state.chat_history = []
            st.rerun()

if __name__ == "__main__":
    main()
